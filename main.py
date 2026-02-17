import os
import sys
import uvicorn
from fastapi import FastAPI, UploadFile, File, Form, HTTPException, Depends, Request
from fastapi.responses import JSONResponse, HTMLResponse
from fastapi.middleware.cors import CORSMiddleware
from fastapi.security.oauth2 import OAuth2PasswordBearer, OAuth2PasswordRequestForm
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles

from pdf2image import convert_from_path
from PIL import Image
import pytesseract
from pytesseract import Output
import re
import uuid
from tinydb import TinyDB, Query
import fitz  # PyMuPDF
import base64
from typing import Optional
import hashlib
from jose import jwt, JWTError
from datetime import datetime, timedelta
from docx import Document
from docx.shared import Pt
from io import BytesIO

def resource_path(relative_path):
    """ Get absolute path to resource, works for dev and for PyInstaller """
    if hasattr(sys, '_MEIPASS'):
        return os.path.join(sys._MEIPASS, relative_path)
    return os.path.join(os.path.abspath("."), relative_path)

app = FastAPI()
oauth2_scheme = OAuth2PasswordBearer(tokenUrl="token")

# --- FIXED MOUNT PATHS ---
static_dir = resource_path("static")
if os.path.isdir(static_dir) and os.listdir(static_dir):
    app.mount("/static", StaticFiles(directory=static_dir), name="static")
templates = Jinja2Templates(directory=resource_path("templates"))

# --- EXTERNAL BINARY CONFIG ---
# We will tell PyInstaller to put all executables in a "bin" folder
pytesseract.pytesseract.tesseract_cmd = resource_path("bin/tesseract") 
os.environ['TESSDATA_PREFIX'] = resource_path("tessdata")

# Critical for pdf2image: It needs to know where the poppler binaries are
poppler_path = resource_path("bin")

# JWT settings
SECRET_KEY = "your-secret-key"
ALGORITHM = "HS256"
ACCESS_TOKEN_EXPIRE_MINUTES = 30
REFRESH_TOKEN_EXPIRE_DAYS = 7

# CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

db = TinyDB("ocr_results.json")
user_db = TinyDB("users.json")
UPLOAD_DIR = "uploads"
os.makedirs(UPLOAD_DIR, exist_ok=True)

# ------------------------ AUTH UTILS ------------------------

def hash_password(password: str) -> str:
    return hashlib.sha256(password.encode()).hexdigest()

def create_access_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(minutes=ACCESS_TOKEN_EXPIRE_MINUTES))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def create_refresh_token(data: dict, expires_delta: Optional[timedelta] = None):
    to_encode = data.copy()
    expire = datetime.utcnow() + (expires_delta or timedelta(days=REFRESH_TOKEN_EXPIRE_DAYS))
    to_encode.update({"exp": expire})
    return jwt.encode(to_encode, SECRET_KEY, algorithm=ALGORITHM)

def verify_token(token: str, credentials_exception):
    try:
        payload = jwt.decode(token, SECRET_KEY, algorithms=[ALGORITHM])
        username: str = payload.get("sub")
        if username is None:
            raise credentials_exception
        return username
    except JWTError:
        raise credentials_exception

async def get_current_user(token: str = Depends(oauth2_scheme)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    return verify_token(token, credentials_exception)

# ------------------------ OCR LOGIC ------------------------

SCRIPT_LANG_MAP = {
    'Latin': 'eng',
    'Meetei_Mayek': 'mni',
    'Devanagari': 'hin',
    'Bengali': 'ben',

}

COMMON_ENG_WORDS = {
    "the", "be", "to", "of", "and", "a", "in", "that", "have", "i", "it", "for", "not", "on", "with", "he", "as", "you", "do", "at",
    "this", "but", "his", "by", "from", "they", "we", "say", "her", "she", "or", "an", "will", "my", "one", "all", "would", "there",
    "their", "what", "so", "up", "out", "if", "about", "who", "get", "which", "go", "me", "when", "make", "can", "like", "time",
    "no", "just", "him", "know", "take", "people", "into", "year", "your", "good", "some", "could", "them", "see", "other", "than",
    "then", "now", "look", "only", "come", "its", "over", "think", "also", "back", "after", "use", "two", "how", "our", "work",
    "first", "well", "way", "even", "new", "want", "because", "any", "these", "give", "day", "most", "us", "writ", "petition",
    "civil", "no", "of", "court", "judgment", "order", "case", "versus", "union", "india", "state", "manipur", "respondents",
    "petitioner", "advocate", "counsel", "honble", "justice", "mr", "mrs", "shri", "smt", "disclaimer", "vernacular", "meant", 
    "restricted", "litigant", "understand", "language", "purpose", "practical", "official", "original", "version", "authentic", 
    "field", "execution", "implementation", "high", "imphal"
}

def pdf_to_images(pdf_path):
    try:
        return convert_from_path(pdf_path, poppler_path=poppler_path)
    except Exception as e:
        raise RuntimeError(f"Failed to convert PDF: {e}")

def preprocess_image(image):
    image = image.convert('L')
    image = image.point(lambda x: 0 if x < 128 else 255, '1')
    return image

def detect_script(image) -> str:
    try:
        osd = pytesseract.image_to_osd(image)
        for line in osd.splitlines():
            if "Script" in line:
                return line.split(":")[-1].strip()
    except Exception:
        pass
    return "Latin"  # Default fallback

def ocr_image(image, lang):
    try:
        available_languages = set(pytesseract.get_languages())
        print(f"Available languages: {available_languages}")
        
        if lang == "auto":
            script = detect_script(image)
            lang = SCRIPT_LANG_MAP.get(script, "eng")
            
        # Filter requested languages
        requested_langs = lang.split('+')
        valid_langs = [l for l in requested_langs if l in available_languages]
        
        print(valid_langs)
        
        if not valid_langs:
            print(f"Warning: No valid languages found in request '{lang}'. Falling back to 'eng' (if available) or first available.")
            if 'eng' in available_languages:
                final_lang = 'eng'
            elif available_languages:
                final_lang = list(available_languages)[0]
            else:
                return None # No languages available at all
        else:
            final_lang = "+".join(valid_langs)
            
        print(f"Performing OCR with language(s): {final_lang}")

        # HYBRID LOGIC for mixed content (specifically mni+eng)
        # Dictionary-based line switching
        if 'mni' in valid_langs and 'eng' in valid_langs:
             print("Using Hybrid Line-Based Dictionary OCR for mixed content...")
             
             # 1. Layout analysis with 'eng' to find lines
             data = pytesseract.image_to_data(image, lang='eng', output_type=Output.DICT)
             
             if 'text' not in data:
                 return pytesseract.image_to_string(image, lang=final_lang)

             n_boxes = len(data['text'])
             lines = {}
             
             # Group by line
             for i in range(n_boxes):
                 if int(data['conf'][i]) == -1: continue
                 
                 # Key: (block, par, line)
                 key = (data['block_num'][i], data['par_num'][i], data['line_num'][i])
                 
                 if key not in lines:
                     lines[key] = {
                         'text': [],
                         'left': [], 'top': [], 'width': [], 'height': []
                     }
                 
                 lines[key]['text'].append(data['text'][i])
                 lines[key]['left'].append(data['left'][i])
                 lines[key]['top'].append(data['top'][i])
                 lines[key]['width'].append(data['width'][i])
                 lines[key]['height'].append(data['height'][i])

             sorted_keys = sorted(lines.keys())
             final_text = ""
             
             for key in sorted_keys:
                 l_data = lines[key]
                 line_text_eng = " ".join(l_data['text']).strip()
                 
                 if not line_text_eng: continue
                 
                 # Check English Score
                 tokens = [re.sub(r'[^a-zA-Z]', '', t).lower() for t in line_text_eng.split()]
                 tokens = [t for t in tokens if t]
                 
                 is_eng = False
                 if tokens:
                     match_count = sum(1 for t in tokens if t in COMMON_ENG_WORDS)
                     score = match_count / len(tokens)
                     if score >= 0.2: # Threshold
                         is_eng = True
                 
                 if is_eng:
                     final_text += line_text_eng + "\n"
                 else:
                     # Re-OCR line with MNI
                     x_min = min(l_data['left'])
                     y_min = min(l_data['top'])
                     x_max = max([l+w for l, w in zip(l_data['left'], l_data['width'])])
                     y_max = max([t+h for t, h in zip(l_data['top'], l_data['height'])])
                     
                     padding = 5
                     crop = image.crop((
                         max(0, x_min - padding),
                         max(0, y_min - padding),
                         min(image.width, x_max + padding),
                         min(image.height, y_max + padding)
                     ))
                     
                     mni_text = pytesseract.image_to_string(crop, lang='mni').strip()
                     final_text += mni_text + "\n"
                     
             return final_text

        return pytesseract.image_to_string(image, lang=final_lang)
    except Exception as e:
        print(f"OCR Error: {e}")
        return None

def extract_text_from_pdf(pdf_path, language, doc_id, original_filename):
    images = pdf_to_images(pdf_path)
    if not images:
        return None

    result = {
        "id": doc_id,
        "source_file": original_filename,
        "language": language,
        "page_count": len(images),
        "pages": []
    }

    for i, image in enumerate(images):
        page_number = i + 1
        processed = preprocess_image(image)
        text = ocr_image(processed, lang=language or "auto")
        print("TEXT:", text)
        result["pages"].append({
            "page_number": page_number,
            "text": text if text else "",
            "status": "success" if text else "ocr_failed"
        })

    return result

# ------------------------ API ROUTES ------------------------

@app.post("/upload/")
async def upload_pdf(file: UploadFile = File(...), lang: str = Form("auto"), current_user: str = Depends(get_current_user)):
    if not file.filename.endswith(".pdf"):
        raise HTTPException(status_code=400, detail="File must be a PDF")

    pdf_path = os.path.join(UPLOAD_DIR, file.filename)
    with open(pdf_path, "wb") as f:
        f.write(await file.read())

    doc_id = str(uuid.uuid4())
    result = extract_text_from_pdf(pdf_path, lang, doc_id, file.filename)
    if result:
        db.insert(result)
        return {"message": "OCR completed", "id": doc_id, "page_count": result["page_count"]}
    else:
        raise HTTPException(status_code=500, detail="OCR failed")

@app.post("/upload_base64/")
async def upload_base64_pdf(b64_string: str = Form(...), filename: Optional[str] = Form("base64_upload.pdf"), lang: str = Form("auto"), current_user: str = Depends(get_current_user)):
    try:
        pdf_data = base64.b64decode(b64_string)
    except Exception:
        raise HTTPException(status_code=400, detail="Invalid base64 string")

    pdf_path = os.path.join(UPLOAD_DIR, filename)
    with open(pdf_path, "wb") as f:
        f.write(pdf_data)

    doc_id = str(uuid.uuid4())
    result = extract_text_from_pdf(pdf_path, lang, doc_id, filename)
    if result:
        db.insert(result)
        return {"message": "OCR completed", "id": doc_id, "page_count": result["page_count"]}
    else:
        raise HTTPException(status_code=500, detail="OCR failed")

@app.get("/search/")
def search_text(query: str, current_user: str = Depends(get_current_user)):
    results = []
    for record in db:
        matched_pages = []
        for page in record["pages"]:
            if query.lower() in page["text"].lower():
                try:
                    pdf_path = os.path.join(UPLOAD_DIR, record["source_file"])
                    doc = fitz.open(pdf_path)
                    page_pix = doc.load_page(page["page_number"] - 1).get_pixmap()
                    img_bytes = page_pix.pil_tobytes(format="PNG")
                    encoded_image = base64.b64encode(img_bytes).decode("utf-8")
                except Exception:
                    encoded_image = None

                matched_pages.append({
                    "page_number": page["page_number"],
                    "text": page["text"],
                    "page_image_base64": encoded_image,
                    "image_format": "image/png"
                })

        if matched_pages:
            results.append({
                "id": record["id"],
                "source_file": record["source_file"],
                "page_count": record["page_count"],
                "matches": matched_pages
            })

    return JSONResponse(content=results)

@app.get("/documents/")
def list_documents(current_user: str = Depends(get_current_user)):
    documents = []
    for record in db.all(): # Use TinyDB for OCR results
        documents.append({
            "id": record["id"],
            "source_file": record["source_file"],
            "page_count": record["page_count"]
        })
    return JSONResponse(content=documents)

@app.get("/document/{doc_id}")
def get_document(doc_id: str, current_user: str = Depends(get_current_user)):
    result = db.get(Query().id == doc_id)
    if result:
        pdf_path = os.path.join(UPLOAD_DIR, result["source_file"])
        if os.path.exists(pdf_path):
            with open(pdf_path, "rb") as f:
                encoded_pdf = base64.b64encode(f.read()).decode("utf-8")
            result["full_pdf_base64"] = encoded_pdf
        return result
    raise HTTPException(status_code=404, detail="Document not found")

@app.get("/document/{doc_id}/edited_docx")
async def get_edited_docx(doc_id: str, current_user: str = Depends(get_current_user)):
    record = db.get(Query().id == doc_id) # Use TinyDB for OCR results
    if not record:
        raise HTTPException(status_code=404, detail="Document not found")

    # Create a DOCX document
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Noto Sans'
    font.size = Pt(12)

    for page_data in record["pages"]:
        text = page_data["text"].replace('\r\n', '\n')
        paragraphs = text.split('\n\n')  # Treat double newlines as paragraph breaks
        for para in paragraphs:
            paragraph = doc.add_paragraph()
            lines = para.split('\n')  # Handle single line breaks
            for i, line in enumerate(lines):
                paragraph.add_run(line)
                if i < len(lines) - 1:
                    paragraph.add_run().add_break()

    # Save to a BytesIO stream
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    # Encode to base64
    encoded_docx = base64.b64encode(buffer.read()).decode("utf-8")
    return JSONResponse(content={
        "edited_docx_base64": encoded_docx,
        "filename": f"edited_{record['source_file'].replace('.pdf', '.docx')}"
    })

@app.get("/document/{doc_id}/page_image/{page_number}")
def get_document_page_image(doc_id: str, page_number: int, current_user: str = Depends(get_current_user)):
    record = db.get(Query().id == doc_id)
    if not record:
        raise HTTPException(status_code=404, detail="Document not found")

    pdf_path = os.path.join(UPLOAD_DIR, record["source_file"])
    if not os.path.exists(pdf_path):
        raise HTTPException(status_code=404, detail="PDF file not found")

    try:
        doc = fitz.open(pdf_path)
        if not (0 < page_number <= doc.page_count):
            raise HTTPException(status_code=400, detail="Page number out of range")

        page = doc.load_page(page_number - 1)
        pix = page.get_pixmap()
        img_bytes = pix.pil_tobytes(format="PNG")
        encoded_image = base64.b64encode(img_bytes).decode("utf-8")
        return JSONResponse(content={"page_image_base64": encoded_image, "image_format": "image/png"})

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error processing PDF page: {e}")

@app.put("/edit/")
def edit_page_text(doc_id: str = Form(...), page_number: int = Form(...), new_text: str = Form(...), current_user: str = Depends(get_current_user)):
    Document = Query()
    document = db.get(Document.id == doc_id)
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")

    updated = False
    for page in document["pages"]:
        if page["page_number"] == page_number:
            page["text"] = new_text
            page["status"] = "edited"
            updated = True
            break

    if not updated:
        raise HTTPException(status_code=404, detail="Page not found in document")

    db.update(document, Document.id == doc_id)
    return {"message": "Text updated", "doc_id": doc_id, "page_number": page_number}

@app.delete("/document/{doc_id}")
def delete_document(doc_id: str, current_user: str = Depends(get_current_user)):
    try:
        # Find the document
        Document = Query()
        document_list = db.search(Document.id == doc_id)
        
        if not document_list:
            raise HTTPException(status_code=404, detail="Document not found")
            
        document = document_list[0]
        
        # Delete the file
        pdf_path = os.path.join(UPLOAD_DIR, document["source_file"])
        if os.path.exists(pdf_path):
            os.remove(pdf_path)
            
        # Delete from database
        db.remove(Document.id == doc_id)
        
        return {"message": "Document deleted successfully", "id": doc_id}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Error deleting document: {str(e)}")

@app.post("/login/")
async def login(form_data: OAuth2PasswordRequestForm = Depends()):
    User = Query()
    hashed_password = hash_password(form_data.password)
    user = user_db.get((User.username == form_data.username) & (User.password == hashed_password))
    if not user:
        raise HTTPException(status_code=401, detail="Invalid username or password")

    access_token = create_access_token(data={"sub": form_data.username})
    refresh_token = create_refresh_token(data={"sub": form_data.username})
    user_db.update({"refresh_token": refresh_token}, User.username == form_data.username)

    return {"access_token": access_token, "refresh_token": refresh_token, "token_type": "bearer"}

@app.post("/refresh/")
async def refresh_token(refresh_token: str = Form(...)):
    credentials_exception = HTTPException(
        status_code=401,
        detail="Could not validate credentials",
        headers={"WWW-Authenticate": "Bearer"},
    )
    refreshed_username = verify_token(refresh_token, credentials_exception)

    User = Query()
    user = user_db.get(User.username == refreshed_username)
    if not user or user.get("refresh_token") != refresh_token:
        raise HTTPException(status_code=401, detail="Invalid refresh token")

    new_access_token = create_access_token(data={"sub": refreshed_username})
    return {"access_token": new_access_token, "token_type": "bearer"}

@app.post("/register/")
def register(username: str = Form(...), password: str = Form(...)):
    User = Query()
    if user_db.get(User.username == username):
        raise HTTPException(status_code=400, detail="Username already exists")
    user_db.insert({"username": username, "password": hash_password(password)})
    return {"message": "User registered", "username": username}


@app.get("/", response_class=HTMLResponse)
async def read_root(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})

@app.get("/login", response_class=HTMLResponse)
async def login_page(request: Request):
    return templates.TemplateResponse("login.html", {"request": request})

@app.get("/register", response_class=HTMLResponse)
async def register_page(request: Request):
    return templates.TemplateResponse("register.html", {"request": request})

@app.get("/dashboard", response_class=HTMLResponse)
async def dashboard_page(request: Request):
    return templates.TemplateResponse("dashboard.html", {"request": request})


if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)